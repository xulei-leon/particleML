from __future__ import annotations

import json
import importlib.util
import re
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
CORPUS = ROOT / "docs" / "references" / "literature-dossier" / "source-corpus.json"
AUDIT = ROOT / "docs" / "references" / "literature-dossier" / "search-audit.md"
BUILDER = ROOT / "docs" / "references" / "literature-dossier" / "build_literature_dossier.py"
OUTPUT = ROOT / "docs" / "references" / "particleml-literature-reference-dossier.docx"
REQUIRED_THEMES = {
    "jet_foundation_models",
    "datasets_and_top_tagging",
    "particle_set_architectures",
    "feature_representation",
    "pretraining_and_transfer",
    "evaluation_and_reproducibility",
    "generation_and_anomaly_detection",
}
VALID_SOURCE_TYPES = {
    "peer_reviewed",
    "preprint",
    "dataset",
    "official_documentation",
}
VALID_MANUSCRIPT_USES = {
    "Introduction",
    "Related Work",
    "Methods",
    "Evaluation",
    "Limitations",
    "Future Work",
}


def load_corpus() -> dict[str, object]:
    return json.loads(CORPUS.read_text(encoding="utf-8"))


def load_builder():
    spec = importlib.util.spec_from_file_location("literature_dossier_builder", BUILDER)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load builder from {BUILDER}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class LiteratureDossierTests(unittest.TestCase):
    def test_verified_corpus_contract(self) -> None:
        corpus = load_corpus()
        metadata = corpus["metadata"]
        sources = corpus["sources"]

        self.assertEqual(metadata["search_date"], "2026-07-15")
        self.assertEqual(metadata["reference_style"], "Group-Report numbered style")
        self.assertGreaterEqual(len(sources), 28)
        self.assertLessEqual(len(sources), 35)
        self.assertEqual({source["theme"] for source in sources}, REQUIRED_THEMES)
        self.assertEqual(len({source["citation_key"] for source in sources}), len(sources))
        self.assertEqual(len({source["title"].casefold() for source in sources}), len(sources))

        dois = [source["doi"].casefold() for source in sources if source.get("doi")]
        self.assertEqual(len(dois), len(set(dois)))

        for source in sources:
            with self.subTest(citation_key=source["citation_key"]):
                self.assertEqual(source["verification_status"], "verified")
                self.assertTrue(source["authors"])
                self.assertTrue(source["title"])
                self.assertGreaterEqual(source["year"], 1980)
                self.assertLessEqual(source["year"], 2026)
                self.assertIn(source["source_type"], VALID_SOURCE_TYPES)
                self.assertTrue(source["verification_url"].startswith("https://"))
                self.assertTrue(source["url"].startswith("https://"))
                self.assertTrue(source["verification_route"])
                self.assertTrue(source["evidence_grade"])
                self.assertTrue(source["project_relevance"])
                self.assertTrue(source["central_contribution"])
                self.assertIn(source["manuscript_use"], VALID_MANUSCRIPT_USES)
                self.assertTrue(source["limitations"])
                if source.get("doi"):
                    self.assertRegex(source["doi"], re.compile(r"10\.\d{4,9}/\S+"))

    def test_search_audit_covers_every_source(self) -> None:
        audit = AUDIT.read_text(encoding="utf-8")
        corpus = load_corpus()

        for source in corpus["sources"]:
            with self.subTest(citation_key=source["citation_key"]):
                self.assertIn(source["citation_key"], audit)

        for heading in (
            "## Search Strategy",
            "## Inclusion and Exclusion Criteria",
            "## Source Verification Log",
            "## Deduplication and Version Decisions",
            "## Excluded Candidates",
            "## Search Limitations",
        ):
            self.assertIn(heading, audit)

    def test_reference_formatter_matches_numbered_style(self) -> None:
        builder = load_builder()
        source = load_corpus()["sources"][0]

        rendered = builder.format_reference(1, source)

        self.assertTrue(rendered.startswith("[1] "))
        self.assertIn(source["title"], rendered)
        if source.get("doi"):
            self.assertIn(f"doi: {source['doi']}", rendered)
        if source.get("arxiv"):
            self.assertIn(f"arXiv: {source['arxiv']}", rendered)
        self.assertIn(f"url: {source['url']}", rendered)

    def test_reference_formatter_does_not_repeat_volume_in_venue(self) -> None:
        builder = load_builder()
        source = next(
            source
            for source in load_corpus()["sources"]
            if source["citation_key"] == "qu2022_particle_transformer"
        )

        rendered = builder.format_reference(7, source)

        self.assertNotIn("PMLR 162 162", rendered)

    def test_reference_formatter_avoids_double_period_after_et_al(self) -> None:
        builder = load_builder()
        source = next(
            source
            for source in load_corpus()["sources"]
            if source["authors"][-1] == "et al."
        )

        rendered = builder.format_reference(6, source)

        self.assertNotIn("et al..", rendered)

    def test_generated_docx_has_required_structure_and_preset(self) -> None:
        from docx import Document
        from docx.shared import Inches, Pt

        document = Document(OUTPUT)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        sources = load_corpus()["sources"]

        for heading in (
            "Search Strategy and Selection",
            "Thematic Roadmap",
            "Annotated Source Catalogue",
            "Coverage and Research Gaps",
            "References",
            "Search Limitations and Disclosure",
        ):
            self.assertIn(heading, text)

        reference_count = sum(
            1 for paragraph in document.paragraphs if re.match(r"^\[\d+\] ", paragraph.text)
        )
        self.assertEqual(reference_count, len(sources))
        self.assertGreaterEqual(len(document.tables), 1)

        section = document.sections[0]
        self.assertEqual(section.page_width, Inches(8.5))
        self.assertEqual(section.page_height, Inches(11))
        self.assertEqual(section.top_margin, Inches(1))
        self.assertEqual(section.right_margin, Inches(1))
        self.assertEqual(section.bottom_margin, Inches(1))
        self.assertEqual(section.left_margin, Inches(1))

        normal = document.styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri")
        self.assertEqual(normal.font.size, Pt(11))
        self.assertAlmostEqual(normal.paragraph_format.line_spacing, 1.25)

    def test_catalogue_is_not_preceded_by_an_explicit_page_break(self) -> None:
        from docx import Document

        document = Document(OUTPUT)
        catalogue_index = next(
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text == "Annotated Source Catalogue"
        )
        preceding = document.paragraphs[catalogue_index - 1]

        self.assertFalse(preceding._p.xpath('.//w:br[@w:type="page"]'))


if __name__ == "__main__":
    unittest.main()
