from __future__ import annotations

import json
import os
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
CORPUS_PATH = ROOT / "docs" / "references" / "literature-dossier" / "source-corpus.json"
OUTPUT_PATH = ROOT / "docs" / "references" / "particleml-literature-reference-dossier.docx"

INK_BLUE = RGBColor(0x0B, 0x25, 0x45)
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5A, 0x65, 0x73)
LIGHT_BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"

THEME_ORDER = [
    "jet_foundation_models",
    "datasets_and_top_tagging",
    "particle_set_architectures",
    "feature_representation",
    "pretraining_and_transfer",
    "evaluation_and_reproducibility",
    "generation_and_anomaly_detection",
]

THEME_LABELS = {
    "jet_foundation_models": "Jet Foundation Models",
    "datasets_and_top_tagging": "Datasets and Top Tagging",
    "particle_set_architectures": "Particle-Set Architectures",
    "feature_representation": "Feature Representation and Physics Priors",
    "pretraining_and_transfer": "Pretraining and Transfer",
    "evaluation_and_reproducibility": "Evaluation and Reproducibility",
    "generation_and_anomaly_detection": "Generation and Anomaly Detection",
}

THEME_ROADMAP = {
    "jet_foundation_models": (
        "Defines the OmniLearn/OmniLearned and OmniJet lineages, multi-task claims, and current pretraining directions.",
        "Introduction; Related Work",
    ),
    "datasets_and_top_tagging": (
        "Establishes JetClass provenance, binary top-tagging context, and community benchmark expectations.",
        "Introduction; Methods",
    ),
    "particle_set_architectures": (
        "Explains permutation-invariant, attention, graph, PET, and Lorentz-equivariant backbones.",
        "Related Work; Methods",
    ),
    "feature_representation": (
        "Connects particle-level information, IRC safety, symmetry, and learned low-level representations.",
        "Introduction; Methods",
    ),
    "pretraining_and_transfer": (
        "Supplies transfer-learning terminology, feature-transfer evidence, and the foundation-model frame.",
        "Introduction; Methods",
    ),
    "evaluation_and_reproducibility": (
        "Supports AUC comparison, dependence-aware uncertainty, repeated runs, and reproducible reporting.",
        "Evaluation; Limitations",
    ),
    "generation_and_anomaly_detection": (
        "Places the tagging study within the wider multi-task jet-foundation-model agenda.",
        "Related Work; Future Work",
    ),
}


def load_corpus(path: Path = CORPUS_PATH) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def validate_corpus(corpus: dict[str, Any]) -> None:
    errors: list[str] = []
    sources = corpus.get("sources", [])
    metadata = corpus.get("metadata", {})

    if metadata.get("search_date") != "2026-07-15":
        errors.append("metadata.search_date must be 2026-07-15")
    if metadata.get("reference_style") != "Group-Report numbered style":
        errors.append("metadata.reference_style is incorrect")
    if not 28 <= len(sources) <= 35:
        errors.append("source count must be between 28 and 35")

    keys = [source.get("citation_key", "") for source in sources]
    titles = [source.get("title", "").casefold() for source in sources]
    dois = [source.get("doi", "").casefold() for source in sources if source.get("doi")]
    if len(keys) != len(set(keys)):
        errors.append("citation keys must be unique")
    if len(titles) != len(set(titles)):
        errors.append("titles must be unique")
    if len(dois) != len(set(dois)):
        errors.append("DOIs must be unique")
    if {source.get("theme") for source in sources} != set(THEME_ORDER):
        errors.append("all seven themes must be represented")

    required = {
        "citation_key",
        "theme",
        "source_type",
        "authors",
        "title",
        "year",
        "venue",
        "url",
        "verification_url",
        "verification_status",
        "verification_route",
        "evidence_grade",
        "project_relevance",
        "central_contribution",
        "manuscript_use",
        "limitations",
    }
    for index, source in enumerate(sources, start=1):
        missing = sorted(key for key in required if not source.get(key))
        if missing:
            errors.append(f"source {index} is missing required values: {', '.join(missing)}")
        if source.get("verification_status") != "verified":
            errors.append(f"source {index} is not verified")
        doi = source.get("doi", "")
        if doi and not re.fullmatch(r"10\.\d{4,9}/\S+", doi):
            errors.append(f"source {index} has an invalid DOI")

    if errors:
        raise ValueError("Invalid source corpus:\n- " + "\n- ".join(errors))


def _display_author(name: str) -> str:
    if name.casefold() == "et al.":
        return "et al."
    if "," not in name:
        return name
    family, given = (part.strip() for part in name.split(",", 1))
    return f"{given} {family}".strip()


def format_authors(authors: list[str]) -> str:
    displayed = [_display_author(author) for author in authors]
    if displayed and displayed[-1] == "et al.":
        return ", ".join(displayed[:-1]) + ", et al."
    if len(displayed) > 6:
        return ", ".join(displayed[:3]) + ", et al."
    if len(displayed) == 1:
        return displayed[0]
    if len(displayed) == 2:
        return f"{displayed[0]} and {displayed[1]}"
    return ", ".join(displayed[:-1]) + f", and {displayed[-1]}"


def format_reference(number: int, source: dict[str, Any]) -> str:
    authors = format_authors(source["authors"])
    author_sentence = authors if authors.endswith(".") else f"{authors}."
    venue = source.get("venue", "")
    details: list[str] = []
    volume = str(source.get("volume", ""))
    if volume and not re.search(rf"(?<!\d){re.escape(volume)}(?!\d)", venue):
        details.append(volume)
    if source.get("issue"):
        details[-1] = f"{details[-1]}.{source['issue']}" if details else str(source["issue"])
    if source.get("pages_or_article"):
        details.append(str(source["pages_or_article"]))

    publication_bits = [venue]
    if details:
        publication_bits.append(", ".join(details))
    publication_bits.append(f"({source['year']})")
    publication = " ".join(bit for bit in publication_bits if bit)

    identifiers: list[str] = []
    if source.get("doi"):
        identifiers.append(f"doi: {source['doi']}")
    if source.get("arxiv"):
        identifiers.append(f"arXiv: {source['arxiv']}")
    identifiers.append(f"url: {source['url']}")
    return (
        f"[{number}] {author_sentence} \"{source['title']}.\" In: {publication}. "
        + ". ".join(identifiers)
        + "."
    )


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, name: str, size: float, color: RGBColor | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, name="Calibri", size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, HEADING_BLUE, 18, 10),
        "Heading 2": (13, HEADING_BLUE, 14, 7),
        "Heading 3": (12, HEADING_DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = document.styles[style_name]
        _set_style_font(style, name="Calibri", size=size, color=color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    _set_style_font(caption, name="Calibri", size=9, color=MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def configure_section(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("particleML  |  Literature Reference Dossier")
    _set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    label = footer_paragraph.add_run("Page ")
    _set_run_font(label, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field.append(field_run)
    footer_paragraph._p.append(field)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_fixed_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("Table column widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_title_page(document: Document, source_count: int) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(54)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("PARTICLEML RESEARCH DOSSIER")
    _set_run_font(run, size=10, color=HEADING_BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Literature Reference Dossier")
    _set_run_font(run, size=26, color=INK_BLUE, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run(
        "Jet foundation models, particle-level feature availability, and JetClass top tagging"
    )
    _set_run_font(run, size=13, color=MUTED)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(28)
    run = metadata.add_run(f"Verified sources: {source_count}   |   Search date: 15 July 2026")
    _set_run_font(run, size=10.5, color=HEADING_DARK_BLUE, bold=True)

    scope = document.add_paragraph()
    scope.paragraph_format.left_indent = Inches(0.45)
    scope.paragraph_format.right_indent = Inches(0.45)
    scope.paragraph_format.space_after = Pt(12)
    scope.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = scope.add_run("Purpose. ")
    _set_run_font(run, size=11, color=INK_BLUE, bold=True)
    run = scope.add_run(
        "This dossier supplies a verified literature base for a publication-oriented study of how "
        "nested particle-level feature sets affect fine-tuning of an OmniLearned/PET-style backbone "
        "on JetClass binary top-versus-QCD tagging. It is a source map and reference catalogue, not "
        "a substitute for reading the cited papers."
    )
    _set_run_font(run, size=11)

    use = document.add_paragraph()
    use.paragraph_format.left_indent = Inches(0.45)
    use.paragraph_format.right_indent = Inches(0.45)
    use.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = use.add_run("How to use it. ")
    _set_run_font(run, size=11, color=INK_BLUE, bold=True)
    run = use.add_run(
        "Start with the thematic roadmap, use the annotations to select evidence for a manuscript "
        "section, then cite the formal numbered reference list after checking the complete source."
    )
    _set_run_font(run, size=11)
    document.add_page_break()


def _add_labeled_paragraph(document: Document, label: str, value: str, *, compact: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3 if compact else 5)
    paragraph.paragraph_format.line_spacing = 1.2 if compact else 1.25
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, size=10.5 if compact else 11, color=INK_BLUE, bold=True)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, size=10.5 if compact else 11)


def _add_search_section(document: Document, source_count: int) -> None:
    document.add_heading("Search Strategy and Selection", level=1)
    paragraph = document.add_paragraph(
        "A targeted scoping search was conducted on 15 July 2026 using project seed sources and "
        "direct HTTPS queries to Crossref, arXiv, INSPIRE-HEP, Zenodo, and official proceedings or "
        "journal pages. The final corpus contains "
    )
    paragraph.add_run(str(source_count)).bold = True
    paragraph.add_run(
        " sources. Every final verification endpoint returned HTTP 200, and every DOI, arXiv "
        "identifier, title, author set, and year was checked against an authoritative record."
    )

    _add_labeled_paragraph(
        document,
        "Inclusion",
        "Direct relevance to a planned paper section; resolvable authoritative metadata; a distinct "
        "scientific role; primary peer-reviewed evidence when available; and explicit preprint labels.",
    )
    _add_labeled_paragraph(
        document,
        "Exclusion",
        "Unresolved or mismatched metadata, duplicate versions, secondary summaries, indirect general "
        "relevance, and sources redundant with stronger project-aligned evidence.",
    )
    _add_labeled_paragraph(
        document,
        "Verification boundary",
        "Bibliographic existence and record-level claims were checked. Full-text reading is not claimed "
        "for every source; core papers must be read before claim-level manuscript citation.",
    )


def _add_roadmap(document: Document, sources: list[dict[str, Any]]) -> None:
    document.add_heading("Thematic Roadmap", level=1)
    document.add_paragraph(
        "The themes below map the verified corpus to the planned paper. Counts refer to each source's "
        "primary theme; several sources can support more than one manuscript section."
    )
    caption = document.add_paragraph("Table 1. Literature coverage by evidence theme.", style="Caption")
    caption.paragraph_format.keep_with_next = True

    counts = Counter(source["theme"] for source in sources)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Evidence theme", "Sources", "Primary contribution", "Manuscript use")
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, LIGHT_BLUE_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        _set_run_font(run, size=9.5, color=INK_BLUE, bold=True)
    _set_repeat_table_header(table.rows[0])

    for row_number, theme in enumerate(THEME_ORDER, start=1):
        focus, use = THEME_ROADMAP[theme]
        row = table.add_row()
        values = (THEME_LABELS[theme], str(counts[theme]), focus, use)
        for cell, value in zip(row.cells, values):
            if row_number % 2 == 0:
                _set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(value)
            _set_run_font(run, size=9.2)

    _set_fixed_table_geometry(table, [1900, 850, 4200, 2410])
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    run = after.add_run(
        "Coverage is strongest for the model lineage, particle-set architectures, and evaluation "
        "methods. The central unresolved question is the controlled effect of downstream feature "
        "availability under a fixed pretrained checkpoint and data protocol."
    )
    _set_run_font(run, size=9.5, color=MUTED, italic=True)


def _source_type_label(source_type: str) -> str:
    return {
        "peer_reviewed": "Peer-reviewed publication",
        "preprint": "Verified preprint or report",
        "dataset": "Authoritative dataset record",
        "official_documentation": "Official documentation",
    }[source_type]


def _add_annotated_catalogue(document: Document, sources: list[dict[str, Any]]) -> None:
    document.add_heading("Annotated Source Catalogue", level=1)
    document.add_paragraph(
        "Entries are grouped by their primary evidence role. Source numbers match the formal reference "
        "list. Annotations separate what a source contributes from how it should be used and from the "
        "limits of that use."
    )

    numbered = {source["citation_key"]: index for index, source in enumerate(sources, start=1)}
    for theme in THEME_ORDER:
        document.add_heading(THEME_LABELS[theme], level=2)
        for source in (item for item in sources if item["theme"] == theme):
            number = numbered[source["citation_key"]]
            heading = document.add_heading(f"Source {number} | {source['title']}", level=3)
            heading.paragraph_format.keep_with_next = True

            citation = document.add_paragraph()
            citation.paragraph_format.space_after = Pt(5)
            citation.paragraph_format.keep_with_next = True
            run = citation.add_run(
                f"{format_authors(source['authors'])} ({source['year']}). {source['venue']}."
            )
            _set_run_font(run, size=10.5, color=MUTED, italic=True)

            _add_labeled_paragraph(
                document,
                "Source and evidence",
                f"{_source_type_label(source['source_type'])}; {source['evidence_grade']}.",
                compact=True,
            )
            _add_labeled_paragraph(
                document,
                "Verification",
                source["verification_route"],
                compact=True,
            )
            _add_labeled_paragraph(
                document,
                "Central contribution",
                source["central_contribution"],
                compact=True,
            )
            _add_labeled_paragraph(
                document,
                "Project relevance",
                source["project_relevance"],
                compact=True,
            )
            _add_labeled_paragraph(
                document,
                "Recommended manuscript use",
                source["manuscript_use"],
                compact=True,
            )
            _add_labeled_paragraph(
                document,
                "Limitation",
                source["limitations"],
                compact=True,
            )


def _add_gap_section(document: Document) -> None:
    document.add_heading("Coverage and Research Gaps", level=1)
    document.add_paragraph(
        "The literature strongly supports the existence of transferable jet representations, the "
        "importance of permutation-aware particle-set models, and the use of JetClass as a large-scale "
        "benchmark. It does not yet answer the project's central causal comparison."
    )
    gaps = [
        (
            "Matched feature-ablation evidence",
            "No retained study isolates nested particle-level feature groups while keeping the raw "
            "JetClass files, split, pretrained checkpoint, training schedule, and evaluation pipeline fixed.",
        ),
        (
            "Checkpoint and pretraining overlap",
            "The paper must document the OmniLearned checkpoint's pretraining corpus and avoid claiming "
            "out-of-domain transfer if JetClass or the same top/QCD task family appears in pretraining.",
        ),
        (
            "Feature-schema validity",
            "Config D remains publication-ready only if every impact-parameter feature is traced to a real "
            "JetClass column and passes leakage, masking, and transformation audits.",
        ),
        (
            "Independent validation",
            "Most model-family evidence is authored by the teams introducing the methods. A simple fixed "
            "Deep Sets/PFN baseline and independent reproduction strengthen feature-ranking claims.",
        ),
        (
            "Uncertainty and dependence",
            "AUC intervals on one test set do not replace repeated-seed uncertainty. Nested data sizes and "
            "shared examples create dependence that must be reflected in comparisons.",
        ),
        (
            "External and multi-task generalization",
            "Generation, anomaly detection, and transfer to other detector domains remain future work until "
            "the core JetClass feature-ablation result is stable."
        ),
    ]
    for label, value in gaps:
        _add_labeled_paragraph(document, label, value)


def _add_references(document: Document, sources: list[dict[str, Any]]) -> None:
    document.add_page_break()
    document.add_heading("References", level=1)
    intro = document.add_paragraph(
        "Numbered format follows the reference presentation in Group-Report.pdf. DOI and arXiv fields are "
        "included only when confirmed in the verified corpus."
    )
    intro.paragraph_format.space_after = Pt(10)
    for number, source in enumerate(sources, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.32)
        paragraph.paragraph_format.first_line_indent = Inches(-0.32)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(format_reference(number, source))
        _set_run_font(run, size=10)


def _add_limitations_and_disclosure(document: Document) -> None:
    document.add_heading("Search Limitations and Disclosure", level=1)
    limitations = [
        "This is a targeted scoping review aligned with the current research plan, not an exhaustive "
        "systematic review or meta-analysis.",
        "Authoritative metadata and abstracts were checked, but record verification is not a claim that "
        "every full text has been read end to end.",
        "No subscription-only citation index was used, and literature published after 15 July 2026 is not "
        "represented.",
        "Recent preprints and new model releases should be refreshed immediately before manuscript submission.",
        "Statistical methods must be adapted to the dependence structure created by shared JetClass splits, "
        "nested training sizes, and repeated stochastic training runs.",
    ]
    for index, value in enumerate(limitations, start=1):
        _add_labeled_paragraph(document, f"Limitation {index}", value, compact=True)

    document.add_heading("AI-Assistance Disclosure", level=2)
    document.add_paragraph(
        "AI-assisted tools were used for literature discovery, metadata normalization, source verification "
        "support, and document formatting. All included references were checked against the authoritative "
        "records identified in the verification audit; inclusion does not imply that every full text has "
        "been read by the researcher."
    )


def build_document(
    corpus_path: Path = CORPUS_PATH,
    output_path: Path = OUTPUT_PATH,
) -> Path:
    corpus = load_corpus(corpus_path)
    validate_corpus(corpus)
    sources = corpus["sources"]

    document = Document()
    configure_styles(document)
    configure_section(document)
    document.core_properties.title = "particleML Literature Reference Dossier"
    document.core_properties.subject = "Verified literature for jet foundation-model research"
    document.core_properties.author = "particleML Research Project"
    document.core_properties.keywords = (
        "jet physics, foundation models, OmniLearned, OmniJet, JetClass, top tagging"
    )

    _add_title_page(document, len(sources))
    _add_search_section(document, len(sources))
    _add_roadmap(document, sources)
    _add_annotated_catalogue(document, sources)
    _add_gap_section(document)
    _add_references(document, sources)
    _add_limitations_and_disclosure(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_suffix(".tmp.docx")
    if temporary.exists():
        temporary.unlink()
    document.save(temporary)
    os.replace(temporary, output_path)
    return output_path


def main() -> None:
    output = build_document()
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
